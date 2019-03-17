import ReadPySource as rps
import GenScreenShot as gss
import time, docx
from docx.shared import Inches,Pt,Mm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT

#print("Input Required Data")
UID="17BCS3527"#input("UID - ")
githandle="@dsp9107"#"@"+input("githandle - ")
aim="WAP to calculate area of a circle"#input("Aim : ")
source="Script.py"#input("Name of File - ")

print("Preparing ...")
time.sleep(0.5)

code=rps.read()

screen=gss.genscreen(source)

print("Working on The Document ...")

doc=docx.Document()

section = doc.sections[0]
section.page_height = Mm(297)
section.page_width = Mm(210)
section.left_margin = Mm(25.4)
section.right_margin = Mm(25.4)
section.top_margin = Mm(25.4)
section.bottom_margin = Mm(25.4)
section.header_distance = Mm(12.5)
section.footer_distance = Mm(12.5)

styles = doc.styles
style = styles.add_style("Head", WD_STYLE_TYPE.PARAGRAPH)
style.base_style = styles["Normal"]
style.font.name = 'Calibri'
tab_stops = style.paragraph_format.tab_stops
tab_stops.add_tab_stop(Inches(3.0), WD_TAB_ALIGNMENT.CENTER)
tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT)

header = section.header
paragraph = header.paragraphs[0]
paragraph.text = "\t\t"+UID
paragraph.style = doc.styles["Head"]

footer = section.footer
paragraph = footer.paragraphs[0]
paragraph.text = "\t\t"+githandle
paragraph.style = doc.styles["Head"]

style=doc.styles.add_style('Heading', WD_STYLE_TYPE.PARAGRAPH)
style.font.name = 'Calibri'
style.font.size = Pt(16)
style.font.bold=True

style=doc.styles.add_style('Text', WD_STYLE_TYPE.PARAGRAPH)
style.font.name = 'Calibri'
style.font.size = Pt(13)

style=doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
style.font.name = 'Courier New'
style.font.size = Pt(10)

p=doc.add_paragraph("Aim :",style='Heading')
p=doc.add_paragraph(aim,style='Text')

p=doc.add_paragraph("Code :",style='Heading')
p=doc.add_paragraph("",style='Code')
for c in code:
    p.add_run(c)

doc.add_paragraph("Output :",style='Heading')
doc.add_picture(screen,width=Inches(5.5))
doc.save(UID+'.docx')
print("Saved !")