import os, time, docx, json
import pyautogui as pyg

from PIL import ImageGrab
from docx.shared import Inches,Pt,Mm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT

def genscreen(filename = "Script.py", path = os.getcwd(), ss = "ScreenShot.png"):
    #Run CMD
    pyg.keyDown('win')
    pyg.press('r')
    pyg.keyUp('win')
    time.sleep(0.5)
    print(".", end = '')

    pyg.typewrite('cmd')
    pyg.press('enter')
    time.sleep(0.5)
    print(".", end = '')
    
    #Change Directory
    pyg.typewrite(path[:2])
    pyg.press('enter')
    time.sleep(0.5)    
    pyg.typewrite('cd ' + path)
    pyg.press('enter')
    time.sleep(0.5)
    print(".", end = '')
    
    #Prepping The Screen
    pyg.typewrite('echo off')
    pyg.press('enter')
    time.sleep(0.5)
    pyg.typewrite('cls')
    pyg.press('enter')
    time.sleep(0.5)
    print(".", end = '')
    
    #Execute Python Script
    pyg.typewrite('python ' + filename)
    pyg.press('enter')
    time.sleep(1)
    print(".", end = '')
    
    #Take Screenshot
    pyg.keyDown('alt')
    pyg.press('prtscr')
    pyg.keyUp('alt')
    time.sleep(1)
    print(".", end = '')
    
    # Grab Image From Clipboard And Save As .png
    im = ImageGrab.grabclipboard()
    im.save("out\\" + ss, 'PNG')
    time.sleep(1)
    print(".", end = '')

    # Exit
    pyg.typewrite('exit')
    pyg.press('enter')
    time.sleep(0.5)
    print(".", end = '')
    return ss

# Read detail.json Into detail
with open("detail.json", "r") as read_file:
    detail = json.load(read_file)

# Initializing Required Data Elements
UID = detail['user']['uid']                 #input("UID - ")
githandle = detail['user']['githandle']     #input("githandle - ")
source = detail['infile']                   #"Script.py"
aim = "WAP to calculate area of a circle"   #input("Aim : ")

print("Preparing")
time.sleep(0.5)

# Create Directory To Output Generated Files
if not os.path.exists('out') :
    os.mkdir('out')
    print(".", end = '')

# Read Code
code = open(source,"r")
print(".", end = '')

# Capture Screenshot
screen = genscreen(source)
print(".", end = '')

doc=docx.Document()

# Set Page Layout
section = doc.sections[0]
section.page_height = Mm(297)
section.page_width = Mm(210)
section.left_margin = Mm(25.4)
section.right_margin = Mm(25.4)
section.top_margin = Mm(25.4)
section.bottom_margin = Mm(25.4)
section.header_distance = Mm(12.5)
section.footer_distance = Mm(12.5)

# Set Page Styles
styles = doc.styles
style = styles.add_style("Head", WD_STYLE_TYPE.PARAGRAPH)
style.base_style = styles["Normal"]
style.font.name = 'Calibri'
tab_stops = style.paragraph_format.tab_stops
tab_stops.add_tab_stop(Inches(3.0), WD_TAB_ALIGNMENT.CENTER)
tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT)

# Add Header
header = section.header
paragraph = header.paragraphs[0]
paragraph.text = "\t\t"+UID
paragraph.style = doc.styles["Head"]

# Add Footer
footer = section.footer
paragraph = footer.paragraphs[0]
paragraph.text = "\t\t@"+githandle
paragraph.style = doc.styles["Head"]

# Set Style - Heading
style = doc.styles.add_style('Heading', WD_STYLE_TYPE.PARAGRAPH)
style.font.name = 'Calibri'
style.font.size = Pt(16)
style.font.bold = True

# Set Style - Text
style = doc.styles.add_style('Text', WD_STYLE_TYPE.PARAGRAPH)
style.font.name = 'Calibri'
style.font.size = Pt(13)

# Set Style - Code
style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
style.font.name = 'Courier New'
style.font.size = Pt(10)

# Add Aim
p = doc.add_paragraph("Aim :", style = 'Heading')
p = doc.add_paragraph(aim, style = 'Text')

# Add Code
p = doc.add_paragraph("Code :", style = 'Heading')
p = doc.add_paragraph("", style = 'Code')
for c in code:
    p.add_run(c)

# Attach Screenshot
doc.add_paragraph("Output :", style = 'Heading')
doc.add_picture('out\\' + screen, width = Inches(5.5))

# Save
doc.save('out\\' + detail['outfile'] + '.docx')
print("\nDone !")
